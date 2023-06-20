import pandas as pd
import numpy as np
import docx

def doc_add_df(doc, df):
  t = doc.add_table(df.shape[0]+1, df.shape[1])

  # add the header rows.
  for j in range(df.shape[-1]):
      t.cell(0,j).text = df.columns[j]

  # add the rest of the data frame
  for i in range(df.shape[0]):
      for j in range(df.shape[-1]):
          t.cell(i+1,j).text = str(df.values[i,j])
          
def add_elements_to_doc(doc, modifiers, headers, texts, kinds, elements, width, heading):
  doc.add_heading(heading, 1)
  for i, el in enumerate(elements):
    if modifiers[i] == 'np':
      doc.add_page_break()
    
    doc.add_heading(headers[i], 3)

    doc.add_paragraph(texts[i])

    if kinds[i]=='t':
      doc_add_df(doc, el)
    
    elif kinds[i]=='f':
      doc.add_picture(el, width=width)
      el.close()
      
def get_text_width(document):
  """
  Returns the text width in mm.
  """
  section = document.sections[0]
  return (section.page_width - section.left_margin - section.right_margin) / 36000
