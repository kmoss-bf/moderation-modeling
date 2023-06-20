import numpy as np
import pandas as pd
import torch
from transformers import BertTokenizer, BertModel, BertConfig
import datetime
from tqdm import tqdm
import os
import transformers
import warnings
from sklearn import metrics
import matplotlib.pyplot as plt
warnings.simplefilter(action='ignore', category=FutureWarning)
import train_model

model_name = 'bert-base-uncased'

config = {
  "lin": 64,
  "dropout": .5,
  "lr": 1e-5,
  "batch_size": 32
}
data_path="/content/drive/MyDrive/data_v1.csv"

bert_model = BertModel.from_pretrained('bert-base-uncased')
tokenizer = BertTokenizer.from_pretrained('bert-base-uncased')

model = train_model.train(config, data_path, bert_model, tokenizer, epochs=1)

df = pd.read_csv(data_path)
new_df = df[['text', 'target']].copy()

escapes = ''.join([chr(char) for char in range(1, 32)])
translator = str.maketrans('', '', escapes)

for ind in tqdm(new_df.index):
    new_df.loc[ind, 'text'] = new_df.loc[ind, 'text'].translate(translator)

dataset = train_model.CustomDataset(new_df, tokenizer, train_model.MAX_LEN)

params = {'batch_size': 32,
                'shuffle': False,
                'num_workers': 0
                }

data_loader = torch.utils.data.DataLoader(dataset, **params)

training_loader, testing_loader = train_model.load_data(data_path=data_path,tokenizer=tokenizer, train_batch_size=config['batch_size'])

device = "cpu"
if torch.cuda.is_available():
  device = "cuda:0"

train_outputs_, train_targets = train_model.validation(training_loader, model=model, device=device)
test_outputs_, test_targets = train_model.validation(testing_loader, model=model, device=device)

data_outputs_, data_targets = train_model.validation(data_loader, model=model, device=device, )

test_outputs_, test_targets = np.array(test_outputs_), np.array(test_targets)

positives = np.histogram(test_outputs_[test_targets==1], range=(0, 1), bins=25)
negatives = np.histogram(test_outputs_[test_targets==0], range=(0, 1), bins=25)

fig_hist = plt.figure(figsize=(12, 6))
plt.bar(np.arange(positives[0].shape[0]), negatives[0]+positives[0], width=.9, color='red')
plt.bar(np.arange(positives[0].shape[0]), positives[0], width=.9, color='green')
plt.xticks(np.arange(positives[0].shape[0]), [f'{(pos*100+2):.0f}' for pos in positives[1][:-1]])
plt.xlabel('Prediction Probability (%)')
plt.ylabel('Frequency')
plt.show()

fpr, tpr, thresholds = metrics.roc_curve(test_targets,test_outputs_)

fig_roc = plt.figure()
plt.scatter(fpr, tpr)
plt.title(f"ROC Curve, AUCROC={metrics.roc_auc_score(test_targets,test_outputs_):.3f}")
plt.show()

test_outputs = (test_outputs_ > 0.5).astype(int)

import seaborn as sns
cm = metrics.confusion_matrix(test_targets, test_outputs)

target_names = ['Not Moderated', 'Moderated']
cmn = cm.astype('float') / cm.sum(axis=1)[:, np.newaxis]
fig_confusion, ax = plt.subplots(figsize=(8,7))
sns.heatmap(cmn, annot=True, fmt='.2f', xticklabels=target_names, yticklabels=target_names)
plt.ylabel('Actual')
plt.xlabel('Predicted')
plt.show(block=False)

precision, recall, threshold = metrics.precision_recall_curve(test_targets, test_outputs_)
# Plot the output.
fig_pr_re = plt.figure()
plt.plot(threshold, precision[:-1], c ='r', label ='PRECISION')
plt.plot(threshold, recall[:-1], c ='b', label ='RECALL')
plt.grid()
plt.legend()
plt.title('Precision-Recall Curve')

df = pd.read_csv('/content/drive/MyDrive/wip2.csv')

df.drop(columns=['Unnamed: 0.1', 'Unnamed: 0'], inplace=True)

df_scores = df[~df['active_score'].isna()]

scores = np.array(df_scores.iloc[:, -3:])

data_outputs = (np.array(data_outputs_)>0.5).astype(float)
data_outputs_scores = data_outputs[(~df['active_score'].isna())]

score_names = ['self-harm', 'depression', 'activity']

fig, ax = plt.subplots(nrows=2, ncols=3, figsize=(15, 8))
for i in range(2):
  for j in range(3):
    sns.histplot(scores[outputs_pretend==i][:, j], ax=ax[i, j])
    ax[i, j].set_title(f'Predicted {i}, {score_names[j]} score')

import docx
import doc_utils
import io

# create an instance of a word document
doc = docx.Document()
# add a heading of level 0 (largest heading)
doc.add_heading('Model Report', 0)
doc.add_heading(f'Model {model_name}')

width = docx.shared.Mm(doc_utils.get_text_width(doc))

figs = [fig_hist, fig_roc, fig_confusion, fig_pr_re, fig]

buffers = [io.BytesIO() for fig_ in figs]
for i, fig_ in enumerate(figs):
  fig_.savefig(buffers[i], format='png', bbox_inches='tight')

for buffer in buffers:
  doc.add_picture(buffer)

doc.add_paragraph(f'{metrics.accuracy_score(test_targets, test_outputs):.3f}, {metrics.f1_score(test_targets, test_outputs):.3f}')

if os.path.isdir(f'/content/drive/MyDrive/models/{model_name}'):
  print("This model already exists. Please change model_type to account for new version.")
else:
  os.mkdir(f'/content/drive/MyDrive/models/{model_name}')
  torch.save(model.state_dict(), f'/content/drive/MyDrive/models/{model_name}/model.pt')
  np.save(f'/content/drive/MyDrive/models/{model_name}/predicted_labels.npy', data_outputs)
  doc.save(f'/content/drive/MyDrive/models/{model_name}/model_description.docx')
